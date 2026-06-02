"""
fill_dokument.py — Galledia Dokument-Skill v1.1
CI-konforme .docx aus Vorlage_Dokument.dotx (V3)

build_document(
    titel, untertitel, datum, rechtseinheit, adresse, empfaenger,
    abschnitte=[{
        'titel': 'Kapitel',
        'inhalt': [
            {'typ': 'text',    'inhalt': 'Fliesstext'},
            {'typ': 'bullet',  'inhalt': 'Aufzählung'},
            {'typ': 'h2',      'inhalt': 'Unterkapitel'},
            {'typ': 'h3',      'inhalt': 'Abschnitt'},
            {'typ': 'tabelle', 'inhalt': [['Kopf A','Kopf B'],['Wert','Wert']]},
        ]
    }],
    output_path='output.docx'
)
"""
import io, os, zipfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_TEMPLATE = os.path.join(os.path.dirname(__file__), "assets", "Vorlage_Dokument.dotx")

# ── Template-Loader ───────────────────────────────────────────────────────────

def _patch_bullet_numbering(numbering_xml_bytes, styles_xml_bytes):
    """Synchronisiert numbering.xml lvl 0 (abstractNum 10 / DokBullet) mit
    den Einzugswerten aus styles.xml DokBullet. Word priorisiert sonst die
    numbering-Werte ueber die Style-Werte; ohne Sync entsteht ein unge-
    wollter Tabstop zwischen Bullet und Text.
    """
    from lxml import etree
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{" + W_NS + "}"

    # Soll-Werte aus styles.xml DokBullet auslesen
    styles = etree.fromstring(styles_xml_bytes)
    left_val, hanging_val = "142", "142"   # Default falls Style fehlt
    for s in styles.findall(f"{W}style"):
        if s.get(f"{W}styleId") != "DokBullet":
            continue
        ind = s.find(f"{W}pPr/{W}ind")
        if ind is not None:
            left_val = ind.get(f"{W}left", left_val)
            hanging_val = ind.get(f"{W}hanging", hanging_val)
        break

    tree = etree.fromstring(numbering_xml_bytes)
    target_abstract_id = None
    for num in tree.findall(f"{W}num"):
        if num.get(f"{W}numId") == "35":
            a = num.find(f"{W}abstractNumId")
            if a is not None:
                target_abstract_id = a.get(f"{W}val")
            break
    if target_abstract_id is None:
        return numbering_xml_bytes
    for an in tree.findall(f"{W}abstractNum"):
        if an.get(f"{W}abstractNumId") != target_abstract_id:
            continue
        lvl0 = an.find(f"{W}lvl[@{W}ilvl='0']")
        if lvl0 is None:
            continue
        pPr = lvl0.find(f"{W}pPr")
        if pPr is None:
            pPr = etree.SubElement(lvl0, f"{W}pPr")
        ind = pPr.find(f"{W}ind")
        if ind is None:
            ind = etree.SubElement(pPr, f"{W}ind")
        ind.set(f"{W}left", left_val)
        ind.set(f"{W}hanging", hanging_val)
        if ind.get(f"{W}firstLine"):
            del ind.attrib[f"{W}firstLine"]
        break
    return etree.tostring(tree, xml_declaration=True, encoding='utf-8', standalone=True)


def _load(path=None):
    """Lädt .dotx als editierbares Document (patcht Content_Types.xml + numbering.xml)."""
    src = path or _TEMPLATE
    # styles.xml einmalig lesen, um DokBullet-Werte fuer numbering-Sync zu kennen
    with zipfile.ZipFile(src, 'r') as zin:
        styles_bytes = zin.read('word/styles.xml')

    buf = io.BytesIO()
    with zipfile.ZipFile(src, 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    data = data.replace(
                        b'.wordprocessingml.template.main+xml',
                        b'.wordprocessingml.document.main+xml')
                elif item.filename == 'word/numbering.xml':
                    data = _patch_bullet_numbering(data, styles_bytes)
                zout.writestr(item, data)
    buf.seek(0)
    return Document(buf)

# ── Primitive ─────────────────────────────────────────────────────────────────

def _force_left_align(para):
    """Erzwingt linksbündige Ausrichtung — überschreibt distribute/justify
    aus Template-Style (z.B. DokAbsender war 'distribute' → "G a l l e d i a")."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)
    # bestehendes <w:jc> entfernen
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)


def _set_text(para, text):
    """Ersetzt Paragraph-Text; \\n-Zeichen → Word-Zeilenumbrüche."""
    for r in para.runs:
        r.text = ''
    lines = str(text).split('\n')
    if para.runs:
        para.runs[0].text = lines[0]
    else:
        para.add_run(lines[0])
    last = para.runs[0]._r
    for line in lines[1:]:
        br = OxmlElement('w:br')
        last.addnext(br); last = br
        r  = OxmlElement('w:r')
        t  = OxmlElement('w:t'); t.text = line; r.append(t)
        last.addnext(r); last = r

def _para(style, text=''):
    """Erstellt ein neues Paragraphen-Element mit Style und Text."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    ps  = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style)
    pPr.append(ps)

    # DokBullet: Numbering-Sync passiert bereits in _load() (numbering.xml wird
    # an Style-Werte angeglichen). Hier keine zusaetzliche Paragraph-Override
    # noetig — Style-Werte aus Vorlage gewinnen.

    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t'); t.text = text
        if text != text.strip():
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t); p.append(r)
    return p

# ── TOC ───────────────────────────────────────────────────────────────────────

def _toc(ref_p):
    """
    Fügt TOC auf eigener Seite ein:
    pageBreakBefore auf der Überschrift → TOC auf neuer Seite.
    Seitenumbruch nach TOC-Feld → Inhalt auf neuer Seite.
    """
    def _fc(ftype):
        r = OxmlElement('w:r'); fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), ftype); r.append(fc); return r

    # TOC-Überschrift mit pageBreakBefore → TOC startet auf eigener Seite
    h = OxmlElement('w:p')
    hPr = OxmlElement('w:pPr')
    hStyle = OxmlElement('w:pStyle')
    hStyle.set(qn('w:val'), 'Inhaltsverzeichnisberschrift')
    hPr.append(hStyle)
    pbb = OxmlElement('w:pageBreakBefore')   # zwingt TOC auf neue Seite
    hPr.append(pbb)
    h.append(hPr)
    hR = OxmlElement('w:r'); hT = OxmlElement('w:t')
    hT.text = 'Inhaltsverzeichnis'; hR.append(hT); h.append(hR)
    ref_p._p.addnext(h)

    # TOC-Feld (Word aktualisiert mit F9)
    tp = OxmlElement('w:p')
    tpPr = OxmlElement('w:pPr'); tpStyle = OxmlElement('w:pStyle')
    tpStyle.set(qn('w:val'), 'Verzeichnis1'); tpPr.append(tpStyle); tp.append(tpPr)
    tp.append(_fc('begin'))
    ri = OxmlElement('w:r'); ins = OxmlElement('w:instrText')
    ins.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '; ri.append(ins); tp.append(ri)
    tp.append(_fc('separate'))
    rt = OxmlElement('w:r'); tt = OxmlElement('w:t')
    tt.text = 'Inhaltsverzeichnis — in Word mit F9 aktualisieren.'
    rt.append(tt); tp.append(rt)
    tp.append(_fc('end'))
    h.addnext(tp)

    # Seitenumbruch nach TOC → Inhalt auf neuer Seite
    pb = OxmlElement('w:p')
    pbR = OxmlElement('w:r'); pbBr = OxmlElement('w:br')
    pbBr.set(qn('w:type'), 'page'); pbR.append(pbBr); pb.append(pbR)
    tp.addnext(pb)

    return pb   # Anchor: Content-Insertion startet nach diesem Element

# ── Tabelle ───────────────────────────────────────────────────────────────────

def _table(rows_data):
    """
    CI-konforme Tabelle.
    Kopfzeile: Volte Semibold, Schwarz, Hintergrund #F2F2F5.
    Datenzeilen: Volte Regular.
    """
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tW = OxmlElement('w:tblW')
    tW.set(qn('w:w'), '5000'); tW.set(qn('w:type'), 'pct')
    tblPr.append(tW); tbl.append(tblPr)

    for ri, row in enumerate(rows_data):
        tr = OxmlElement('w:tr')
        for cell in row:
            tc  = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            if ri == 0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'F2F2F5')
                tcPr.append(shd)
            tc.append(tcPr)

            cp  = OxmlElement('w:p')
            cr  = OxmlElement('w:r')
            crPr = OxmlElement('w:rPr')

            # Font: Volte Semibold (Kopf) / Volte Regular (Daten)
            fn = OxmlElement('w:rFonts')
            fn.set(qn('w:ascii'), 'Volte Semibold' if ri == 0 else 'Volte')
            fn.set(qn('w:hAnsi'), 'Volte Semibold' if ri == 0 else 'Volte')
            crPr.append(fn)

            # Kopfzeile: Schwarz, kein Fett-Tag (Semibold reicht)
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '000000')  # immer Schwarz
            crPr.append(col)

            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); crPr.append(sz)

            cr.append(crPr)
            ct = OxmlElement('w:t'); ct.text = str(cell)
            cr.append(ct); cp.append(cr); tc.append(cp); tr.append(tc)
        tbl.append(tr)
    return tbl

# ── Hauptfunktion ─────────────────────────────────────────────────────────────

def build_document(
    titel,
    untertitel    = '',
    datum         = '',
    rechtseinheit = '',
    adresse       = '',    # \n für Zeilenumbrüche: 'Strasse\nPLZ Ort'
    empfaenger    = '',    # optional, \n für Zeilenumbrüche
    abschnitte    = None,
    output_path   = 'output.docx',
    template      = None
):
    """
    Erstellt ein CI-konformes Galledia-Dokument aus Vorlage_Dokument.dotx.

    Aufbau:  Deckblatt | Inhaltsverzeichnis (eigene Seite) | Inhalt
    Schluss: Ort/Datum und Rechtseinheit stehen auf dem Deckblatt — nicht wiederholt.

    abschnitte = [{'titel': str, 'inhalt': [{'typ': str, 'inhalt': str|list}]}]
    typ-Werte:  'text' | 'bullet' | 'h2' | 'h3' | 'tabelle'
    """
    doc = _load(template)

    # ── 1. Leere Inhaltsverzeichnis-Überschriften entfernen (Vorlage-Artefakte) ─
    for p in list(doc.paragraphs):
        if 'Inhaltsverzeichnis' in p.style.name and not p.text.strip():
            p._p.getparent().remove(p._p)

    # ── 2. Deckblatt-Platzhalter befüllen ────────────────────────────────────
    datum_set = False
    rechtseinheit_set = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if   t == '[Dokumenttitel]':
            _set_text(p, titel)
        elif t == '[Untertitel / Anlass]':
            _set_text(p, untertitel)
        elif '[Ort], [Datum]' in t and not datum_set:
            _set_text(p, datum); datum_set = True
        elif '[Rechtseinheit] / [Adresse]' in t and not rechtseinheit_set:
            _set_text(p, rechtseinheit + ('\n' + adresse if adresse else ''))
            _force_left_align(p)   # Fix Blocksatz "G a l l e d i a"
            rechtseinheit_set = True
        elif t == '[Empfänger-Block, optional]':
            if empfaenger:
                _set_text(p, empfaenger)
            else:
                # Kein Empfänger → Verfasser-Block per Spacer nach unten
                # schieben (sonst klebt er direkt unter dem Datum).
                _set_text(p, "")
                last = p._p
                for _ in range(34):
                    sp = _para('Standard', ''); last.addnext(sp); last = sp

    # ── 3. Fusszeile: [Dokumenttitel] → Titel ────────────────────────────────
    for sect in doc.sections:
        for hf in (sect.footer, getattr(sect, 'first_page_footer', None)):
            if hf is None:
                continue
            for fp in hf.paragraphs:
                if '[Dokumenttitel]' in fp.text:
                    for run in fp.runs:
                        run.text = run.text.replace('[Dokumenttitel]', titel)

    # ── 4. TOC einfügen (eigene Seite) ───────────────────────────────────────
    toc_end = None
    for p in doc.paragraphs:
        if p.text.strip() == '[Inhaltsverzeichnis]':
            toc_end = _toc(p)
            p._p.getparent().remove(p._p)
            break

    # ── 5. Inhalt einfügen ───────────────────────────────────────────────────
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == '[Text / Inhalt]':
            anchor = p; break

    if anchor and abschnitte:
        last = anchor._p
        for si, section in enumerate(abschnitte):
            if si > 0:   # Leerzeile zwischen Kapiteln
                sp = _para('Standard', ''); last.addnext(sp); last = sp
            h1 = _para('berschrift1', section.get('titel', ''))
            last.addnext(h1); last = h1

            items = section.get('inhalt', [])
            style_map = {
                'text':   'Standard',
                'bullet': 'DokBullet',
                'h2':     'berschrift2',
                'h3':     'berschrift3',
            }
            for i, item in enumerate(items):
                typ = item.get('typ', 'text')
                cnt = item.get('inhalt', '')
                nxt = items[i+1].get('typ', 'text') if i+1 < len(items) else None

                if typ in ('h2', 'h3'):   # Leerzeile vor Unterüberschrift
                    sp = _para('Standard', ''); last.addnext(sp); last = sp

                if typ == 'tabelle':
                    tbl = _table(cnt)
                    last.addnext(tbl); last = tbl
                    sp  = _para('Standard', ''); last.addnext(sp); last = sp
                else:
                    np = _para(style_map.get(typ, 'Standard'), cnt)
                    last.addnext(np); last = np

                # Leerzeile nach: text immer; bullet nur wenn naechster Block
                # kein weiteres bullet ist (Listen bleiben kompakt)
                needs_spacer = (
                    typ == 'text' and nxt not in (None, 'h2', 'h3')
                    or typ == 'bullet' and nxt not in (None, 'bullet')
                )
                if needs_spacer:
                    sp = _para('Standard', ''); last.addnext(sp); last = sp

        anchor._p.getparent().remove(anchor._p)

    # ── 6. Trailing-Platzhalter am Doc-Ende entfernen ────────────────────────
    # Vorlage hat einen Schluss-Block "[Ort], [Datum] / [Rechtseinheit]" — nach
    # Befuellung Deckblatt bleibt der zweite Treffer als Artefakt. Entfernen.
    _PLACEHOLDER_RESTS = (
        '[Ort], [Datum]',
        '[Rechtseinheit] / [Adresse]',
        '[Rechtseinheit]',
        '[Dokumenttitel]',
        '[Untertitel / Anlass]',
        '[Empfänger-Block, optional]',
        '[Text / Inhalt]',
        '[Inhaltsverzeichnis]',
    )
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if any(ph in t for ph in _PLACEHOLDER_RESTS):
            p._p.getparent().remove(p._p)

    doc.save(output_path)
    return output_path
